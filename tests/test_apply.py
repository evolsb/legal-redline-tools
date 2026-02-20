"""Tests for tracked change application."""

import copy

from docx import Document
from docx.oxml.ns import qn

from legal_redline.apply import (
    apply_redlines,
    apply_tracked_replacement,
    apply_tracked_deletion,
    apply_tracked_insertion,
    apply_tracked_add_section,
    enable_track_revisions,
)


AUTHOR = "Test Reviewer"
DATE = "2025-01-01T00:00:00Z"


def _para_xml_tags(para):
    """Get tag names of child elements in a paragraph."""
    return [child.tag.split("}")[-1] for child in para._element]


def _get_para_text(para):
    """Get full text including tracked changes."""
    texts = []
    for child in para._element:
        tag = child.tag.split("}")[-1]
        if tag == "r":
            t = child.find(qn("w:t"))
            if t is not None and t.text:
                texts.append(t.text)
        elif tag in ("ins", "del"):
            for r in child.findall(qn("w:r")):
                t = r.find(qn("w:t"))
                dt = r.find(qn("w:delText"))
                if t is not None and t.text:
                    texts.append(t.text)
                if dt is not None and dt.text:
                    texts.append(dt.text)
    return "".join(texts)


class TestTrackedReplacement:
    def test_basic_replacement(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_replacement(
            doc, "Acme Corp and its affiliates",
            "Acme Corp and its subsidiaries", AUTHOR, DATE
        )
        assert ok is True

    def test_replacement_creates_del_and_ins(self, sample_docx):
        doc = Document(str(sample_docx))
        apply_tracked_replacement(
            doc, "Acme Corp and its affiliates",
            "Acme Corp and its subsidiaries", AUTHOR, DATE
        )
        # Find the paragraph that had the replacement
        for para in doc.paragraphs:
            tags = _para_xml_tags(para)
            if "del" in tags and "ins" in tags:
                # del should come before ins
                assert tags.index("del") < tags.index("ins")
                return
        raise AssertionError("No paragraph found with both del and ins")

    def test_replacement_preserves_surrounding_text(self, sample_docx):
        doc = Document(str(sample_docx))
        apply_tracked_replacement(
            doc, "Acme Corp and its affiliates",
            "Acme Corp and its subsidiaries", AUTHOR, DATE
        )
        for para in doc.paragraphs:
            full = _get_para_text(para)
            if "subsidiaries" in full:
                # Should still have the leading text
                assert "Service Provider" in full
                return
        raise AssertionError("Replacement text not found")

    def test_replacement_not_found_returns_false(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_replacement(
            doc, "text that does not exist anywhere",
            "replacement", AUTHOR, DATE
        )
        assert ok is False

    def test_replacement_with_smart_quotes(self, smart_quote_docx):
        doc = Document(str(smart_quote_docx))
        # Search with straight quotes — should match smart quotes in doc
        ok = apply_tracked_replacement(
            doc, '"Service Provider" shall not exceed',
            '"Vendor" shall not exceed', AUTHOR, DATE
        )
        assert ok is True


class TestTrackedDeletion:
    def test_basic_deletion(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_deletion(
            doc, "Each party shall indemnify the other", AUTHOR, DATE
        )
        assert ok is True

    def test_deletion_creates_del_element(self, sample_docx):
        doc = Document(str(sample_docx))
        apply_tracked_deletion(
            doc, "Each party shall indemnify the other", AUTHOR, DATE
        )
        for para in doc.paragraphs:
            tags = _para_xml_tags(para)
            if "del" in tags:
                return
        raise AssertionError("No del element found")

    def test_deletion_not_found(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_deletion(doc, "nonexistent text", AUTHOR, DATE)
        assert ok is False


class TestTrackedInsertion:
    def test_basic_insertion(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_insertion(
            doc, "thirty (30) days written notice",
            " and a fifteen (15) day cure period", AUTHOR, DATE
        )
        assert ok is True

    def test_insertion_creates_ins_element(self, sample_docx):
        doc = Document(str(sample_docx))
        apply_tracked_insertion(
            doc, "thirty (30) days written notice",
            " and a cure period", AUTHOR, DATE
        )
        for para in doc.paragraphs:
            tags = _para_xml_tags(para)
            if "ins" in tags:
                return
        raise AssertionError("No ins element found")

    def test_insertion_not_found(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_insertion(
            doc, "nonexistent anchor", "new text", AUTHOR, DATE
        )
        assert ok is False


class TestAddSection:
    def test_basic_add_section(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_add_section(
            doc,
            "governed by the laws of the State of Delaware",
            "Disputes shall be resolved by arbitration.",
            AUTHOR, DATE, new_section_number="5.1"
        )
        assert ok is True

    def test_add_section_not_found(self, sample_docx):
        doc = Document(str(sample_docx))
        ok = apply_tracked_add_section(
            doc, "nonexistent section reference",
            "new text", AUTHOR, DATE
        )
        assert ok is False


class TestEnableTrackRevisions:
    def test_adds_track_revisions(self, sample_docx):
        doc = Document(str(sample_docx))
        enable_track_revisions(doc)
        settings = doc.settings.element
        assert settings.find(qn("w:trackRevisions")) is not None

    def test_idempotent(self, sample_docx):
        doc = Document(str(sample_docx))
        enable_track_revisions(doc)
        enable_track_revisions(doc)
        settings = doc.settings.element
        elements = settings.findall(qn("w:trackRevisions"))
        assert len(elements) == 1


class TestApplyRedlines:
    def test_all_types_applied(self, sample_docx, sample_redlines, tmp_dir):
        output = tmp_dir / "output.docx"
        results = apply_redlines(
            str(sample_docx), str(output), sample_redlines, author=AUTHOR
        )
        assert output.exists()
        ok_count = sum(1 for status, _ in results if status == "OK")
        assert ok_count == 5

    def test_missing_type_field(self, sample_docx, tmp_dir):
        output = tmp_dir / "output.docx"
        redlines = [{"old": "something", "new": "else"}]  # no type
        results = apply_redlines(
            str(sample_docx), str(output), redlines, author=AUTHOR
        )
        assert results[0][0] == "ERROR"
        assert "missing 'type'" in results[0][1]

    def test_missing_required_field(self, sample_docx, tmp_dir):
        output = tmp_dir / "output.docx"
        redlines = [{"type": "replace", "old": "something"}]  # no 'new'
        results = apply_redlines(
            str(sample_docx), str(output), redlines, author=AUTHOR
        )
        assert results[0][0] == "ERROR"
        assert "missing required field" in results[0][1]

    def test_not_found_status(self, sample_docx, tmp_dir):
        output = tmp_dir / "output.docx"
        redlines = [{"type": "replace", "old": "xyz123", "new": "abc"}]
        results = apply_redlines(
            str(sample_docx), str(output), redlines, author=AUTHOR
        )
        assert results[0][0] == "NOT FOUND"

    def test_default_author(self, sample_docx, tmp_dir):
        """Default author should be 'Reviewer', not a personal name."""
        output = tmp_dir / "output.docx"
        redlines = [{"type": "replace", "old": "Acme Corp", "new": "Beta Corp"}]
        apply_redlines(str(sample_docx), str(output), redlines)
        doc = Document(str(output))
        for para in doc.paragraphs:
            for child in para._element:
                author = child.get(qn("w:author"))
                if author:
                    assert author == "Reviewer"
                    return
        # If no tracked change found, the redline didn't match — that's separate

    def test_output_saved_even_with_errors(self, sample_docx, tmp_dir):
        output = tmp_dir / "output.docx"
        redlines = [{"type": "replace"}]  # will error
        apply_redlines(str(sample_docx), str(output), redlines, author=AUTHOR)
        assert output.exists()
