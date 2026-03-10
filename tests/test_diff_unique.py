"""Tests for diff uniqueness expansion (ensure_unique feature)."""

import json

from docx import Document

from legal_redline.diff import (
    _count_occurrences,
    _expand_to_unique,
    _expand_replacement_new,
    _ensure_unique_anchors,
    _build_full_doc_text,
    diff_documents,
)


class TestCountOccurrences:
    def test_no_match(self):
        assert _count_occurrences("hello world", "xyz") == 0

    def test_single_match(self):
        assert _count_occurrences("hello world", "hello") == 1

    def test_multiple_matches(self):
        assert _count_occurrences("the cat and the dog and the bird", "the") == 3

    def test_empty_search(self):
        assert _count_occurrences("hello", "") == 0

    def test_overlapping(self):
        # "aa" in "aaa" — non-overlapping would be 1, but our impl counts 2
        assert _count_occurrences("aaa", "aa") == 2


class TestExpandToUnique:
    def test_already_unique(self):
        source = "The liability cap is twelve months."
        full = "Section 1. The liability cap is twelve months.\nSection 2. Termination clause."
        result = _expand_to_unique("twelve months", source, full)
        assert result == "twelve months"

    def test_expands_ambiguous_snippet(self):
        source_para = "set forth in the applicable Order Form and subject to the terms"
        full = (
            "as defined in the applicable Order Form below.\n"
            "set forth in the applicable Order Form and subject to the terms\n"
            "referenced in the applicable Order Form appendix."
        )
        result = _expand_to_unique("applicable Order Form", source_para, full)
        # Should have expanded to include surrounding context
        assert "applicable Order Form" in result
        assert len(result) > len("applicable Order Form")
        assert _count_occurrences(full, result) == 1

    def test_respects_max_len(self):
        source_para = "A " * 200 + "unique_marker " + "A " * 200
        full = source_para + "\n" + "A " * 50
        result = _expand_to_unique("A", source_para, full, max_len=50)
        assert len(result) <= 60  # some slack for word boundary snapping

    def test_empty_snippet(self):
        assert _expand_to_unique("", "hello world", "hello world") == ""

    def test_empty_source_para(self):
        assert _expand_to_unique("hello", "", "hello world") == "hello"

    def test_snippet_not_in_source_para(self):
        result = _expand_to_unique("xyz", "hello world", "hello world xyz hello world")
        assert result == "xyz"

    def test_snaps_to_word_boundaries(self):
        source_para = "The applicable Order Form governs all terms herein."
        full = (
            "See applicable Order Form for details.\n"
            "The applicable Order Form governs all terms herein."
        )
        result = _expand_to_unique("applicable Order Form", source_para, full)
        # Should not cut words — should snap to word boundaries
        words = result.split()
        # First and last words should be complete (no partial words)
        assert all(c.isalpha() or c in ".,;:!?'\"()-" for c in words[0])
        assert all(c.isalpha() or c in ".,;:!?'\"()-" for c in words[-1])


class TestExpandReplacementNew:
    def test_basic_expansion(self):
        result = _expand_replacement_new(
            original_old="applicable Order Form",
            expanded_old="in the applicable Order Form and",
            original_new="applicable Statement of Work",
            source_para="set forth in the applicable Order Form and subject to terms",
        )
        assert result == "in the applicable Statement of Work and"

    def test_no_expansion_needed(self):
        result = _expand_replacement_new(
            original_old="twelve months",
            expanded_old="twelve months",
            original_new="twenty-four months",
            source_para="liability cap of twelve months preceding the claim",
        )
        assert result == "twenty-four months"

    def test_prefix_only(self):
        result = _expand_replacement_new(
            original_old="Order Form",
            expanded_old="applicable Order Form",
            original_new="Statement of Work",
            source_para="the applicable Order Form governs",
        )
        assert result == "applicable Statement of Work"

    def test_suffix_only(self):
        result = _expand_replacement_new(
            original_old="Order Form",
            expanded_old="Order Form governs",
            original_new="Statement of Work",
            source_para="the applicable Order Form governs",
        )
        assert result == "Statement of Work governs"


class TestEnsureUniqueAnchors:
    def test_replace_expanded(self):
        old_paras = [
            "See the applicable Order Form for pricing.",
            "Terms set forth in the applicable Order Form and conditions apply.",
        ]
        redlines = [
            {
                "type": "replace",
                "old": "applicable Order Form",
                "new": "applicable SOW",
                "title": "Change 1",
                "_source_para": old_paras[1],
            }
        ]
        _ensure_unique_anchors(redlines, old_paras)
        # "old" should have been expanded
        assert _count_occurrences(
            _build_full_doc_text(old_paras), redlines[0]["old"]
        ) == 1
        # "new" should carry the same context
        assert "applicable SOW" in redlines[0]["new"]
        # _source_para should be stripped
        assert "_source_para" not in redlines[0]

    def test_delete_expanded(self):
        old_paras = [
            "Party A shall provide thirty (30) days notice.",
            "Party B shall provide thirty (30) days notice and a cure period.",
        ]
        redlines = [
            {
                "type": "delete",
                "text": "thirty (30) days",
                "title": "Change 1",
                "_source_para": old_paras[1],
            }
        ]
        _ensure_unique_anchors(redlines, old_paras)
        assert _count_occurrences(
            _build_full_doc_text(old_paras), redlines[0]["text"]
        ) == 1

    def test_insert_after_anchor_expanded(self):
        old_paras = [
            "Subject to the terms of this Agreement.",
            "Governed by the terms of this Agreement and applicable law.",
        ]
        redlines = [
            {
                "type": "insert_after",
                "anchor": "terms of this Agreement",
                "text": " (as amended)",
                "title": "Change 1",
                "_source_para": old_paras[1],
            }
        ]
        _ensure_unique_anchors(redlines, old_paras)
        assert _count_occurrences(
            _build_full_doc_text(old_paras), redlines[0]["anchor"]
        ) == 1

    def test_already_unique_not_modified(self):
        old_paras = [
            "The liability cap is twelve months.",
            "Termination requires thirty days notice.",
        ]
        redlines = [
            {
                "type": "replace",
                "old": "twelve months",
                "new": "twenty-four months",
                "title": "Change 1",
                "_source_para": old_paras[0],
            }
        ]
        _ensure_unique_anchors(redlines, old_paras)
        assert redlines[0]["old"] == "twelve months"
        assert redlines[0]["new"] == "twenty-four months"

    def test_no_source_para_skipped(self):
        """Redlines without _source_para (e.g., hand-written) are left alone."""
        redlines = [
            {
                "type": "replace",
                "old": "some text",
                "new": "other text",
                "title": "Change 1",
            }
        ]
        _ensure_unique_anchors(redlines, ["some text appears twice. some text again."])
        assert redlines[0]["old"] == "some text"


class TestDiffDocumentsEnsureUnique:
    """Integration tests: diff two docx files and verify uniqueness."""

    def _make_docx(self, tmp_path, name, paragraphs):
        path = tmp_path / name
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(str(path))
        return path

    def test_diff_produces_unique_old_fields(self, tmp_path):
        """The key scenario: repeated phrase changed in only one location."""
        old_paras = [
            "1.1 The fees are set forth in the applicable Order Form.",
            "4.1 Pricing is per the applicable Order Form schedule.",
            "4.2 Changes to the applicable Order Form require written consent.",
        ]
        # Change only 4.2
        new_paras = [
            "1.1 The fees are set forth in the applicable Order Form.",
            "4.1 Pricing is per the applicable Order Form schedule.",
            "4.2 Changes to the applicable Statement of Work require written consent.",
        ]
        old_path = self._make_docx(tmp_path, "old.docx", old_paras)
        new_path = self._make_docx(tmp_path, "new.docx", new_paras)

        redlines = diff_documents(str(old_path), str(new_path), ensure_unique=True)

        assert len(redlines) >= 1
        full_text = _build_full_doc_text(old_paras)
        for r in redlines:
            if r["type"] == "replace":
                assert _count_occurrences(full_text, r["old"]) == 1, (
                    f"'old' field not unique: '{r['old']}'"
                )

    def test_diff_no_ensure_unique_may_be_ambiguous(self, tmp_path):
        """With ensure_unique=False, short snippets may be ambiguous."""
        old_paras = [
            "1.1 The fees are set forth in the applicable Order Form.",
            "4.1 Pricing is per the applicable Order Form schedule.",
            "4.2 Changes to the applicable Order Form require written consent.",
        ]
        new_paras = [
            "1.1 The fees are set forth in the applicable Order Form.",
            "4.1 Pricing is per the applicable Order Form schedule.",
            "4.2 Changes to the applicable Statement of Work require written consent.",
        ]
        old_path = self._make_docx(tmp_path, "old.docx", old_paras)
        new_path = self._make_docx(tmp_path, "new.docx", new_paras)

        redlines = diff_documents(str(old_path), str(new_path), ensure_unique=False)

        assert len(redlines) >= 1
        # With ensure_unique=False, the raw "old" might be "applicable Order Form"
        # which appears 3 times — this is the bug the feature fixes
        full_text = _build_full_doc_text(old_paras)
        for r in redlines:
            if r["type"] == "replace":
                # Just verify it has the changed text (no uniqueness guarantee)
                assert "Order Form" in r["old"] or "Statement" in r["new"]

    def test_diff_unique_with_deletions(self, tmp_path):
        """Delete redlines should also have unique text fields."""
        old_paras = [
            "The Provider shall comply with all applicable laws.",
            "The Customer shall comply with all applicable laws and regulations.",
        ]
        new_paras = [
            "The Provider shall comply with all applicable laws.",
            "The Customer shall comply with regulations.",
        ]
        old_path = self._make_docx(tmp_path, "old.docx", old_paras)
        new_path = self._make_docx(tmp_path, "new.docx", new_paras)

        redlines = diff_documents(str(old_path), str(new_path), ensure_unique=True)
        full_text = _build_full_doc_text(old_paras)

        for r in redlines:
            if r["type"] == "delete":
                assert _count_occurrences(full_text, r["text"]) == 1, (
                    f"'text' field not unique: '{r['text']}'"
                )

    def test_no_source_para_in_output(self, tmp_path):
        """Internal _source_para metadata should never appear in output."""
        old_paras = ["Hello world.", "Goodbye world."]
        new_paras = ["Hello universe.", "Goodbye world."]
        old_path = self._make_docx(tmp_path, "old.docx", old_paras)
        new_path = self._make_docx(tmp_path, "new.docx", new_paras)

        for ensure in (True, False):
            redlines = diff_documents(
                str(old_path), str(new_path), ensure_unique=ensure,
            )
            for r in redlines:
                assert "_source_para" not in r

    def test_roundtrip_diff_apply(self, tmp_path):
        """Diff + apply should produce correct tracked changes at the right location."""
        from legal_redline.apply import apply_redlines

        old_paras = [
            "1.1 The fees are set forth in the applicable Order Form.",
            "4.1 Pricing is per the applicable Order Form schedule.",
            "4.2 Changes to the applicable Order Form require written consent.",
        ]
        new_paras = [
            "1.1 The fees are set forth in the applicable Order Form.",
            "4.1 Pricing is per the applicable Order Form schedule.",
            "4.2 Changes to the applicable Statement of Work require written consent.",
        ]
        old_path = self._make_docx(tmp_path, "old.docx", old_paras)
        new_path = self._make_docx(tmp_path, "new.docx", new_paras)
        output_path = tmp_path / "tracked.docx"

        redlines = diff_documents(str(old_path), str(new_path), ensure_unique=True)
        results = apply_redlines(str(old_path), str(output_path), redlines)

        # All redlines should apply successfully
        ok_count = sum(1 for status, _ in results if status == "OK")
        assert ok_count == len(redlines), (
            f"Only {ok_count}/{len(redlines)} applied: {results}"
        )

        # Verify the change landed in the right paragraph (4.2, not 1.1)
        doc = Document(str(output_path))
        para_texts = [p.text for p in doc.paragraphs]
        # Para 0 (1.1) should be unchanged — no tracked changes
        from docx.oxml.ns import qn
        para_0 = doc.paragraphs[0]._element
        assert para_0.find(qn("w:del")) is None, (
            "Paragraph 1.1 should NOT have tracked changes"
        )
        assert para_0.find(qn("w:ins")) is None, (
            "Paragraph 1.1 should NOT have tracked changes"
        )
