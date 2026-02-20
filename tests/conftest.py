"""Shared fixtures for legal-redline-tools tests."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def tmp_dir(tmp_path):
    """Provide a temporary directory."""
    return tmp_path


@pytest.fixture
def sample_docx(tmp_path):
    """Create a simple contract .docx for testing."""
    path = tmp_path / "contract.docx"
    doc = Document()
    doc.add_paragraph("1. Definitions")
    doc.add_paragraph(
        'The term "Service Provider" means Acme Corp and its affiliates.'
    )
    doc.add_paragraph("2. Liability")
    doc.add_paragraph(
        "The total aggregate liability of Service Provider shall not exceed "
        "the fees paid in the twelve (12) months preceding the claim."
    )
    doc.add_paragraph("3. Indemnification")
    doc.add_paragraph(
        "Each party shall indemnify the other against all losses, damages, "
        "and expenses arising from a breach of this Agreement."
    )
    doc.add_paragraph("4. Termination")
    doc.add_paragraph(
        "Either party may terminate this Agreement upon thirty (30) days "
        "written notice to the other party."
    )
    doc.add_paragraph("5. Governing Law")
    doc.add_paragraph(
        "This Agreement shall be governed by the laws of the State of Delaware."
    )
    doc.save(str(path))
    return path


@pytest.fixture
def smart_quote_docx(tmp_path):
    """Create a .docx with smart quotes and em dashes (common in legal docs)."""
    path = tmp_path / "smart_quotes.docx"
    doc = Document()
    doc.add_paragraph(
        "The \u201cService Provider\u201d shall not exceed the Buyer\u2019s "
        "expectations \u2014 including performance targets."
    )
    doc.save(str(path))
    return path


@pytest.fixture
def multi_run_docx(tmp_path):
    """Create a .docx where text is split across multiple runs."""
    path = tmp_path / "multi_run.docx"
    doc = Document()
    para = doc.add_paragraph()
    run1 = para.add_run("The total liability ")
    run1.bold = True
    run2 = para.add_run("shall not exceed ")
    run3 = para.add_run("one million dollars.")
    run3.italic = True
    doc.save(str(path))
    return path


@pytest.fixture
def sample_redlines():
    """Standard set of redlines covering all 4 types."""
    return [
        {
            "type": "replace",
            "section": "1",
            "title": "Service Provider Definition",
            "old": "Acme Corp and its affiliates",
            "new": "Acme Corp and its subsidiaries",
        },
        {
            "type": "replace",
            "section": "2",
            "title": "Liability Cap",
            "old": "twelve (12) months",
            "new": "twenty-four (24) months",
        },
        {
            "type": "delete",
            "section": "3",
            "title": "Mutual Indemnification",
            "text": "Each party shall indemnify the other against all losses, damages, and expenses arising from a breach of this Agreement.",
        },
        {
            "type": "insert_after",
            "section": "4",
            "title": "Cure Period",
            "anchor": "thirty (30) days written notice",
            "text": " and a fifteen (15) day cure period",
        },
        {
            "type": "add_section",
            "after_section": "This Agreement shall be governed by the laws of the State of Delaware.",
            "new_section_number": "5.1",
            "text": "Any disputes shall be resolved by binding arbitration.",
        },
    ]
