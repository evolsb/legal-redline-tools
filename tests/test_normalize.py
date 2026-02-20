"""Tests for text normalization and fuzzy matching."""

from legal_redline.apply import _normalize_text, _contains_normalized, _find_text_across_runs
from docx import Document


class TestNormalizeText:
    def test_smart_single_quotes(self):
        assert _normalize_text("\u2018hello\u2019") == "'hello'"

    def test_smart_double_quotes(self):
        assert _normalize_text("\u201Chello\u201D") == '"hello"'

    def test_en_dash(self):
        assert _normalize_text("a\u2013b") == "a-b"

    def test_em_dash(self):
        assert _normalize_text("a\u2014b") == "a-b"

    def test_multiple_spaces(self):
        assert _normalize_text("hello   world") == "hello world"

    def test_tabs_to_single_space(self):
        assert _normalize_text("hello\tworld") == "hello world"

    def test_mixed_whitespace(self):
        assert _normalize_text("hello \t  world") == "hello world"

    def test_no_change_needed(self):
        text = "plain text with normal quotes"
        assert _normalize_text(text) == text

    def test_combined_normalization(self):
        text = "\u201CHello\u201D \u2014  world\u2019s"
        assert _normalize_text(text) == '"Hello" - world\'s'


class TestContainsNormalized:
    def test_exact_match(self):
        assert _contains_normalized("hello world", "hello")

    def test_smart_quote_match(self):
        assert _contains_normalized(
            'The \u201cService Provider\u201d shall',
            'The "Service Provider" shall'
        )

    def test_whitespace_match(self):
        assert _contains_normalized("hello   world", "hello world")

    def test_no_match(self):
        assert not _contains_normalized("hello world", "goodbye")


class TestFindTextAcrossRuns:
    def test_single_run_exact(self):
        doc = Document()
        para = doc.add_paragraph("The total liability shall not exceed one million.")
        result = _find_text_across_runs(para, "liability shall not")
        assert result is not None
        _, _, _, _, matched = result
        assert matched == "liability shall not"

    def test_across_multiple_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("The total ")
        para.add_run("liability shall ")
        para.add_run("not exceed.")
        result = _find_text_across_runs(para, "liability shall not")
        assert result is not None
        start_run, _, end_run, _, matched = result
        assert matched == "liability shall not"
        # Spans from run index 1 into run index 2
        assert start_run == 1
        assert end_run == 2

    def test_smart_quote_fallback(self):
        doc = Document()
        para = doc.add_paragraph('The \u201cService Provider\u201d shall comply.')
        result = _find_text_across_runs(para, 'The "Service Provider" shall')
        assert result is not None

    def test_not_found(self):
        doc = Document()
        para = doc.add_paragraph("Hello world")
        result = _find_text_across_runs(para, "nonexistent text")
        assert result is None

    def test_empty_paragraph(self):
        doc = Document()
        para = doc.add_paragraph("")
        result = _find_text_across_runs(para, "anything")
        assert result is None
