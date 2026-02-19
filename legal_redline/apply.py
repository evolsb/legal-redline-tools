"""Apply tracked changes (redlines) to Word documents via OOXML manipulation."""

import copy
import json
from datetime import datetime, timezone

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_rev_id_counter = 100000


def next_rev_id():
    global _rev_id_counter
    _rev_id_counter += 1
    return str(_rev_id_counter)


def enable_track_revisions(doc):
    """Set the trackRevisions flag in document settings."""
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        track_rev = OxmlElement("w:trackRevisions")
        settings.append(track_rev)


def _make_ins(text, author, date_str, rpr=None):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), next_rev_id())
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _make_del(text, author, date_str, rpr=None):
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), next_rev_id())
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    del_elem.append(r)
    return del_elem


def _make_trailing_run(text, rpr=None):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _get_full_paragraph_text(para):
    return "".join(run.text or "" for run in para.runs)


def _find_text_across_runs(para, search_text):
    """
    Find search_text across potentially split runs.
    Returns (start_run_idx, start_offset, end_run_idx, end_offset) or None.
    """
    runs = para.runs
    if not runs:
        return None

    char_map = []
    for i, run in enumerate(runs):
        for j in range(len(run.text or "")):
            char_map.append((i, j))

    full_text = _get_full_paragraph_text(para)
    pos = full_text.find(search_text)
    if pos < 0:
        return None

    start_run_idx, start_offset = char_map[pos]
    end_pos = pos + len(search_text) - 1
    end_run_idx, end_offset = char_map[end_pos]
    return (start_run_idx, start_offset, end_run_idx, end_offset)


def apply_tracked_replacement(doc, search_text, replacement_text, author, date_str):
    """Replace search_text with replacement_text as tracked changes."""
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if search_text not in full_text:
            continue

        result = _find_text_across_runs(para, search_text)
        if result is None:
            continue

        start_run_idx, start_offset, end_run_idx, end_offset = result
        runs = para.runs
        parent = para._element

        first_rpr = runs[start_run_idx]._element.find(qn("w:rPr"))

        end_run_text = runs[end_run_idx].text or ""
        trailing_text = end_run_text[end_offset + 1:]

        start_run_text = runs[start_run_idx].text or ""
        leading_text = start_run_text[:start_offset]

        all_children = list(parent)
        insert_pos = all_children.index(runs[start_run_idx]._element)

        for i in range(start_run_idx, end_run_idx + 1):
            parent.remove(runs[i]._element)

        new_elements = []
        if leading_text:
            new_elements.append(_make_trailing_run(leading_text, first_rpr))
        new_elements.append(_make_del(search_text, author, date_str, first_rpr))
        new_elements.append(_make_ins(replacement_text, author, date_str, first_rpr))
        if trailing_text:
            new_elements.append(_make_trailing_run(trailing_text, first_rpr))

        for i, elem in enumerate(new_elements):
            parent.insert(insert_pos + i, elem)

        return True
    return False


def apply_tracked_deletion(doc, delete_text, author, date_str):
    """Delete text as a tracked deletion."""
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if delete_text not in full_text:
            continue

        result = _find_text_across_runs(para, delete_text)
        if result is None:
            continue

        start_run_idx, start_offset, end_run_idx, end_offset = result
        runs = para.runs
        parent = para._element

        first_rpr = runs[start_run_idx]._element.find(qn("w:rPr"))

        end_run_text = (runs[end_run_idx].text or "")
        trailing_text = end_run_text[end_offset + 1:]

        start_run_text = (runs[start_run_idx].text or "")
        leading_text = start_run_text[:start_offset]

        insert_pos = list(parent).index(runs[start_run_idx]._element)

        for i in range(start_run_idx, end_run_idx + 1):
            parent.remove(runs[i]._element)

        new_elements = []
        if leading_text:
            new_elements.append(_make_trailing_run(leading_text, first_rpr))
        new_elements.append(_make_del(delete_text, author, date_str, first_rpr))
        if trailing_text:
            new_elements.append(_make_trailing_run(trailing_text, first_rpr))

        for i, elem in enumerate(new_elements):
            parent.insert(insert_pos + i, elem)

        return True
    return False


def apply_tracked_insertion(doc, after_text, new_text, author, date_str):
    """Insert new text after anchor text as a tracked insertion."""
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if after_text not in full_text:
            continue

        result = _find_text_across_runs(para, after_text)
        if result is None:
            continue

        _, _, end_run_idx, end_offset = result
        runs = para.runs
        parent = para._element

        end_run = runs[end_run_idx]
        end_rpr = end_run._element.find(qn("w:rPr"))
        end_run_text = end_run.text or ""

        trailing_text = end_run_text[end_offset + 1:]
        if trailing_text:
            end_run.text = end_run_text[:end_offset + 1]

        insert_pos = list(parent).index(end_run._element) + 1

        ins_elem = _make_ins(new_text, author, date_str, end_rpr)
        parent.insert(insert_pos, ins_elem)

        if trailing_text:
            parent.insert(insert_pos + 1, _make_trailing_run(trailing_text, end_rpr))

        return True
    return False


def apply_redlines(input_path, output_path, redlines, author="Chris Sheehan"):
    """
    Apply a list of redlines to a .docx file.

    Args:
        input_path: Path to original .docx
        output_path: Path for output .docx with tracked changes
        redlines: List of dicts with type, old/new/text/anchor fields
        author: Author name for tracked changes

    Returns:
        List of (status, description) tuples
    """
    doc = Document(input_path)
    enable_track_revisions(doc)

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    results = []
    for i, redline in enumerate(redlines):
        rtype = redline["type"]
        success = False

        if rtype == "replace":
            old = redline["old"]
            new = redline["new"]
            success = apply_tracked_replacement(doc, old, new, author, date_str)
            desc = f"Replace: '{old[:50]}...' -> '{new[:50]}...'"

        elif rtype == "delete":
            text = redline["text"]
            success = apply_tracked_deletion(doc, text, author, date_str)
            desc = f"Delete: '{text[:60]}...'"

        elif rtype == "insert_after":
            anchor = redline["anchor"]
            text = redline["text"]
            success = apply_tracked_insertion(doc, anchor, text, author, date_str)
            desc = f"Insert after: '{anchor[:40]}...'"

        else:
            desc = f"Unknown type: {rtype}"

        status = "OK" if success else "NOT FOUND"
        results.append((status, desc))
        print(f"  [{status}] {desc}")

    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Applied: {sum(1 for s, _ in results if s == 'OK')}/{len(results)} redlines")
    return results
