"""
Render a full-document redline PDF from an original .docx and a list of changes.

Produces a lawyer-style document with:
  - Full document text rendered
  - Red strikethrough for deletions
  - Blue underline for insertions
  - Change bars in the left margin
  - Header on every page
  - Summary table at the end
"""

import html
import re
from datetime import datetime

from docx import Document
from fpdf import FPDF


# ── Colors ──

RED_HEX = "#C81E1E"
BLUE_HEX = "#1E3CB4"
BLACK_HEX = "#000000"
GRAY_HEX = "#808080"

RED = (200, 30, 30)
BLUE = (30, 60, 180)
BLACK = (0, 0, 0)
GRAY = (128, 128, 128)
LIGHT_GRAY = (200, 200, 200)


def _sanitize(text):
    """Replace non-latin1 characters and escape HTML entities."""
    replacements = {
        "\u2014": "--", "\u2013": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2026": "...", "\u2022": "*",
        "\u00a0": " ", "\u200b": "", "\u2003": "  ", "\u2002": " ",
        "\u00b7": "*",
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _esc(text):
    """Sanitize and HTML-escape text for write_html."""
    return html.escape(_sanitize(text))


class RedlinePDF(FPDF):
    """PDF renderer with redline support, headers, footers, and change bars."""

    def __init__(self, header_text=None, **kwargs):
        super().__init__(**kwargs)
        self._header_text = header_text
        self._change_bar_ranges = []  # [(page, y_start, y_end), ...]

    def header(self):
        if self._header_text:
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(*GRAY)
            self.cell(0, 8, _sanitize(self._header_text), align="R")
            self.ln(3)
            self.set_draw_color(*LIGHT_GRAY)
            self.set_line_width(0.2)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(*GRAY)
        self.cell(0, 10, f"Page {self.page_no()}", align="R")

    def draw_change_bars(self):
        """Draw all change bars. Call after document is complete."""
        for page, y_start, y_end in self._change_bar_ranges:
            self.page = page
            self.set_draw_color(0, 0, 0)
            self.set_line_width(0.7)
            x = self.l_margin - 5
            self.line(x, y_start, x, y_end)

    def add_change_bar(self, y_start, y_end):
        """Record a change bar for the current page."""
        if y_end > y_start + 1:
            self._change_bar_ranges.append((self.page, y_start, y_end))


def _get_paragraph_text(para):
    """Get full text of a paragraph."""
    return "".join(run.text or "" for run in para.runs)


def _detect_heading_level(para):
    """Detect if paragraph is a heading and return level (1-6) or 0 for body."""
    style_name = (para.style.name or "").lower()
    if "heading" in style_name:
        for i in range(1, 7):
            if str(i) in style_name:
                return i
        return 1
    return 0


def _is_bold_paragraph(para):
    """Check if the entire paragraph is bold (common in contracts for defined terms)."""
    if not para.runs:
        return False
    return all(run.bold for run in para.runs if run.text and run.text.strip())


def _build_redline_segments(full_text, redlines):
    """
    Given a paragraph's full text and the list of redlines, identify which
    redlines apply and build a list of segments for rendering.

    Returns list of dicts:
        {"text": str, "type": "normal"|"deleted"|"inserted"}

    Also returns the list of applied redline indices.
    """
    # Find all redlines that match this paragraph
    matches = []
    for idx, rl in enumerate(redlines):
        rtype = rl["type"]
        if rtype == "replace":
            pos = full_text.find(rl["old"])
            if pos >= 0:
                matches.append((pos, pos + len(rl["old"]), idx, rtype))
        elif rtype == "delete":
            pos = full_text.find(rl["text"])
            if pos >= 0:
                matches.append((pos, pos + len(rl["text"]), idx, rtype))
        elif rtype == "insert_after":
            pos = full_text.find(rl["anchor"])
            if pos >= 0:
                end = pos + len(rl["anchor"])
                matches.append((end, end, idx, rtype))  # zero-width match at insertion point

    if not matches:
        return [{"text": full_text, "type": "normal"}], []

    # Sort by position (leftmost first)
    matches.sort(key=lambda m: (m[0], m[1]))

    segments = []
    applied = []
    cursor = 0

    for match_start, match_end, rl_idx, rtype in matches:
        rl = redlines[rl_idx]

        # Text before this match
        if match_start > cursor:
            segments.append({"text": full_text[cursor:match_start], "type": "normal"})

        if rtype == "replace":
            segments.append({"text": rl["old"], "type": "deleted"})
            segments.append({"text": rl["new"], "type": "inserted"})
            cursor = match_end
        elif rtype == "delete":
            segments.append({"text": rl["text"], "type": "deleted"})
            cursor = match_end
        elif rtype == "insert_after":
            # The anchor text was already included as normal text above
            # (or will be — we need to include it)
            # Actually, anchor ends at match_start==match_end, so cursor was at match_start
            # The anchor text itself is normal, insertion goes after it
            segments.append({"text": rl["text"], "type": "inserted"})
            cursor = match_end  # same as match_start

        applied.append(rl_idx)

    # Remaining text
    if cursor < len(full_text):
        segments.append({"text": full_text[cursor:], "type": "normal"})

    return segments, applied


def _segments_to_html(segments, para_bold=False):
    """Convert segments to an HTML string for fpdf2's write_html."""
    parts = []
    for seg in segments:
        text = _esc(seg["text"])
        if not text:
            continue

        if seg["type"] == "deleted":
            parts.append(f'<font color="{RED_HEX}"><s>{text}</s></font>')
        elif seg["type"] == "inserted":
            parts.append(f'<font color="{BLUE_HEX}"><u>{text}</u></font>')
        else:
            if para_bold:
                parts.append(f"<b>{text}</b>")
            else:
                parts.append(text)

    return "".join(parts)


def render_redline_pdf(docx_path, redlines, pdf_path, header_text=None,
                       author=None, date_str=None):
    """
    Render a full-document redline PDF.

    Args:
        docx_path: Path to the original .docx file
        redlines: List of redline dicts (same format as apply_redlines)
        pdf_path: Output PDF path
        header_text: Text for page header (e.g. "Proposed Redlines - Feb 2026")
        author: Author name (for cover info)
        date_str: Date string (for cover info)
    """
    doc = Document(docx_path)

    if date_str is None:
        date_str = datetime.now().strftime("%B %d, %Y")

    pdf = RedlinePDF(header_text=header_text)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(20)
    pdf.set_right_margin(15)
    pdf.add_page()

    # Track which redlines were applied (for summary)
    applied_redlines = set()
    change_count = 0
    changes_by_type = {"replace": 0, "delete": 0, "insert_after": 0}

    # ── Render each paragraph ──
    line_height = 4.5
    body_size = 9
    total_paras = len(doc.paragraphs)

    for para_idx, para in enumerate(doc.paragraphs):
        full_text = _get_paragraph_text(para)

        # Skip truly empty paragraphs but add spacing
        if not full_text.strip():
            pdf.ln(3)
            continue

        heading_level = _detect_heading_level(para)
        para_bold = _is_bold_paragraph(para)

        # ── Heading formatting ──
        if heading_level > 0:
            sizes = {1: 14, 2: 12, 3: 11, 4: 10, 5: 10, 6: 9}
            font_size = sizes.get(heading_level, 10)
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", font_size)
            pdf.set_text_color(*BLACK)
            pdf.multi_cell(0, font_size * 0.55, _sanitize(full_text))
            pdf.ln(2)
            continue

        # ── Body paragraph: check for redlines ──
        segments, applied = _build_redline_segments(full_text, redlines)
        has_changes = len(applied) > 0

        if has_changes:
            for rl_idx in applied:
                applied_redlines.add(rl_idx)
                rtype = redlines[rl_idx]["type"]
                if rtype in changes_by_type:
                    changes_by_type[rtype] += 1
                change_count += 1

        # Record Y position for change bar
        y_start = pdf.get_y()

        # Build HTML and render
        para_html = _segments_to_html(segments, para_bold=para_bold)

        if para_html.strip():
            pdf.set_font("Helvetica", "", body_size)
            pdf.set_text_color(*BLACK)
            pdf.write_html(para_html)
            pdf.ln(line_height)

        y_end = pdf.get_y()

        # Add change bar
        if has_changes:
            pdf.add_change_bar(y_start, y_end)

    # ── Summary page ──
    _add_summary_page(pdf, redlines, applied_redlines, changes_by_type,
                      change_count, author, date_str, total_paras)

    # Draw all change bars (must be done after all pages are created)
    current_page = pdf.page
    pdf.draw_change_bars()
    pdf.page = current_page

    pdf.output(pdf_path)
    print(f"Full redline PDF: {pdf_path}")
    print(f"Changes rendered: {change_count}")
    return change_count



def _add_summary_page(pdf, redlines, applied_redlines, changes_by_type,
                      change_count, author, date_str, total_paras):
    """Add a summary/legend page at the end of the redline PDF."""
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(*BLACK)
    pdf.cell(0, 10, "REDLINE SUMMARY", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    # Legend
    pdf.set_draw_color(*LIGHT_GRAY)
    pdf.set_line_width(0.3)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(0, 6, "Legend", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    legend_items = [
        (RED, "Red strikethrough", "Deleted text"),
        (BLUE, "Blue underline", "Inserted text"),
        (BLACK, "Black", "Unchanged text"),
    ]
    for color, style_desc, meaning in legend_items:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*color)
        pdf.cell(40, 5, style_desc, new_x="END")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, meaning, new_x="LMARGIN", new_y="NEXT")

    # Change bar legend
    pdf.ln(1)
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(*BLACK)
    x = pdf.get_x()
    y = pdf.get_y()
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.7)
    pdf.line(x + 2, y, x + 2, y + 5)
    pdf.set_x(x + 8)
    pdf.cell(32, 5, "Change bar", new_x="END")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 5, "Marks paragraphs containing changes", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(4)
    pdf.set_draw_color(*LIGHT_GRAY)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(5)

    # Statistics
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*BLACK)
    pdf.cell(0, 6, "Statistics", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    stats = [
        ("Total changes", str(change_count)),
        ("Replacements", str(changes_by_type.get("replace", 0))),
        ("Deletions", str(changes_by_type.get("delete", 0))),
        ("Insertions", str(changes_by_type.get("insert_after", 0))),
        ("Document paragraphs", str(total_paras)),
    ]
    if author:
        stats.insert(0, ("Author", author))
    if date_str:
        stats.insert(1, ("Date", date_str))

    for label, value in stats:
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*GRAY)
        pdf.cell(40, 5, label + ":", new_x="END")
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*BLACK)
        pdf.cell(0, 5, value, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(5)
    pdf.set_draw_color(*LIGHT_GRAY)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(5)

    # Change table
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*BLACK)
    pdf.cell(0, 6, "Changes", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # Render each change as a bordered block (not a cramped table)
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin

    for idx, rl in enumerate(redlines):
        rtype = rl["type"]
        applied = idx in applied_redlines
        section = rl.get("section", "")
        title = rl.get("title", "")

        if rtype == "replace":
            type_label = "Replacement"
            original = _sanitize(rl["old"])
            proposed = _sanitize(rl["new"])
        elif rtype == "delete":
            type_label = "Deletion"
            original = _sanitize(rl["text"])
            proposed = None
        elif rtype == "insert_after":
            type_label = "Insertion"
            original = _sanitize(f"After: {rl['anchor']}")
            proposed = _sanitize(rl["text"])
        else:
            type_label = rtype
            original = ""
            proposed = None

        # Check if we need a new page
        if pdf.get_y() > pdf.h - 45:
            pdf.add_page()

        # Change header line
        label_parts = [f"{idx + 1}."]
        if section:
            label_parts.append(f"Section {section}")
        if title:
            label_parts.append(f"-- {title}")
        label_parts.append(f"({type_label})")
        if not applied:
            label_parts.append("[NOT FOUND]")

        pdf.set_font("Helvetica", "B", 8)
        hdr_color = BLACK if applied else RED
        pdf.set_text_color(*hdr_color)
        pdf.cell(0, 5, _sanitize(" ".join(label_parts)),
                 new_x="LMARGIN", new_y="NEXT")

        # Original text
        indent = pdf.l_margin + 4
        pdf.set_x(indent)
        pdf.set_font("Courier", "", 7)
        color = RED if rtype in ("replace", "delete") else BLACK
        pdf.set_text_color(*color)
        trunc = original[:120] + ("..." if len(original) > 120 else "")
        pdf.multi_cell(usable_w - 8, 3.5, trunc, new_x="LMARGIN", new_y="NEXT")

        # Proposed text
        if proposed is not None:
            pdf.set_x(indent)
            pdf.set_font("Helvetica", "", 7)
            pdf.set_text_color(*GRAY)
            pdf.cell(8, 3.5, "->", new_x="END")
            pdf.set_font("Courier", "", 7)
            pdf.set_text_color(*BLUE)
            trunc = proposed[:120] + ("..." if len(proposed) > 120 else "")
            pdf.multi_cell(usable_w - 16, 3.5, trunc, new_x="LMARGIN", new_y="NEXT")

        pdf.ln(2)
        pdf.set_draw_color(*LIGHT_GRAY)
        pdf.line(pdf.l_margin + 5, pdf.get_y(), pdf.w - pdf.r_margin - 5, pdf.get_y())
        pdf.ln(2)

    # Footer
    pdf.ln(8)
    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 4, "Generated by legal-redline-tools",
             new_x="LMARGIN", new_y="NEXT", align="C")
